# Script pour créer un template d'attestation de travail
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Créer un nouveau document
    doc = Document()
    
    # Configuration de la page
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Header de l'entreprise
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("ENTREPRISE XYZ")
    header_run.bold = True
    header_run.font.size = Inches(0.2)
    
    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_header.add_run("Service des Ressources Humaines")
    
    # Espacement
    doc.add_paragraph()
    
    # Titre principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ATTESTATION DE TRAVAIL")
    title_run.bold = True
    title_run.font.size = Inches(0.25)
    title_run.underline = True
    
    # Espacement
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Corps de l'attestation
    doc.add_paragraph("Je soussigné(e), [DIRECTEUR_RH], Directeur des Ressources Humaines de [NOM_ENTREPRISE],")
    
    doc.add_paragraph()
    
    certify_para = doc.add_paragraph()
    certify_run = certify_para.add_run("ATTESTE PAR LA PRÉSENTE que :")
    certify_run.bold = True
    
    doc.add_paragraph()
    
    # Informations de l'employé
    doc.add_paragraph("Monsieur/Madame : [NOM_COMPLET]")
    doc.add_paragraph("Né(e) le : [DATE_NAISSANCE]")
    doc.add_paragraph("À : [LIEU_NAISSANCE]")
    
    doc.add_paragraph()
    
    # Informations d'emploi
    employment_para = doc.add_paragraph("A été employé(e) dans notre entreprise en qualité de ")
    employment_para.add_run("[POSTE]").bold = True
    employment_para.add_run(" depuis le ")
    employment_para.add_run("[DATE_ENTREE]").bold = True
    employment_para.add_run(".")
    
    doc.add_paragraph()
    
    # Formule de politesse
    doc.add_paragraph("Cette attestation est délivrée à l'intéressé(e) pour servir et valoir ce que de droit.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date et lieu
    doc.add_paragraph("Fait à [VILLE], le [DATE_GENERATION]")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    signature_para = doc.add_paragraph()
    signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_para.add_run("Le Directeur des Ressources Humaines")
    
    doc.add_paragraph()
    
    signature_name = doc.add_paragraph()
    signature_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_name.add_run("[SIGNATURE]")
    
    # Sauvegarder le template
    doc.save('attestation de travail.docx')
    print("✅ Template 'attestation de travail.docx' créé avec succès!")
    
except ImportError:
    print("❌ Erreur: python-docx n'est pas installé.")
    print("Pour installer: pip install python-docx")
except Exception as e:
    print(f"❌ Erreur lors de la création du template: {e}")
