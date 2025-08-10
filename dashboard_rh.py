import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import numpy as np
import warnings
warnings.filterwarnings('ignore')

# Imports pour la génération d'attestations
try:
    from docx import Document
    from docx.shared import Inches
    import io
    import os
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configuration de la page Streamlit
st.set_page_config(
    page_title="Dashboard RH Executive | Analytics & Insights",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personnalisé pour un design plus professionnel
st.markdown("""
<style>
    .main > div {
        padding-top: 2rem;
    }
    .stMetric {
        background-color: #f8f9fa;
        border: 1px solid #e9ecef;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #007bff;
    }
    .metric-container {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 10px;
        color: white;
        margin: 10px 0;
    }
    .sidebar .sidebar-content {
        background: linear-gradient(180deg, #f8f9fa 0%, #e9ecef 100%);
    }
    .reportview-container .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    h1 {
        color: #2c3e50;
        font-family: 'Arial', sans-serif;
        border-bottom: 3px solid #3498db;
        padding-bottom: 10px;
    }
    h2, h3 {
        color: #34495e;
        font-family: 'Arial', sans-serif;
    }
    .highlight-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #ff6b6b;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #28a745;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #d1ecf1;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 5px solid #17a2b8;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Fonction pour charger et nettoyer les données
@st.cache_data
def load_and_clean_data():
    """Charge et nettoie les données RH du fichier CSV"""
    try:
        # Chargement du fichier CSV avec le bon séparateur
        df = pd.read_csv('Book1.csv', sep=';', encoding='latin-1')
        
        # Suppression des colonnes vides à la fin
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
        df = df.dropna(how='all', axis=1)
        
        # Nettoyage des colonnes vides supplémentaires
        empty_cols = [col for col in df.columns if df[col].isna().all() or (df[col] == '').all()]
        df = df.drop(columns=empty_cols)
        
        # Suppression des lignes vides
        df = df.dropna(how='all')
        
        # Nettoyage des noms de colonnes
        df.columns = df.columns.str.strip()
        
        # Conversion des dates
        if 'Date de naissance' in df.columns:
            df['Date de naissance'] = pd.to_datetime(df['Date de naissance'], format='%d/%m/%Y', errors='coerce')
        
        if 'DateEntree' in df.columns:
            df['DateEntree'] = pd.to_datetime(df['DateEntree'], format='%d/%m/%Y', errors='coerce')
        
        # Calcul de l'âge actuel basé sur la date de naissance
        if 'Date de naissance' in df.columns:
            today = datetime.now()
            df['Age_calcule'] = ((today - df['Date de naissance']).dt.days / 365.25).round().astype('Int64')
        
        # Calcul de l'ancienneté en années
        if 'DateEntree' in df.columns:
            df['Anciennete_calculee'] = ((datetime.now() - df['DateEntree']).dt.days / 365.25).round(1)
        
        # Nettoyage des espaces dans les colonnes textuelles
        text_columns = ['Sexe', 'Situation Civile', 'Type de contrat', 'Direction', 'Déparetement', 'CSP']
        for col in text_columns:
            if col in df.columns:
                df[col] = df[col].astype(str).str.strip()
        
        # Standardisation des valeurs
        if 'Sexe' in df.columns:
            df['Sexe'] = df['Sexe'].map({'M': 'Masculin', 'F': 'Féminin'}).fillna(df['Sexe'])
        
        # Création de catégories d'analyse
        if 'Age_calcule' in df.columns:
            df['Generation'] = pd.cut(df['Age_calcule'], 
                                    bins=[0, 30, 40, 50, 60, 100], 
                                    labels=['Gen Z/Y', 'Millennials', 'Gen X', 'Baby Boomers', 'Seniors'])
            
            df['Statut_Retraite'] = df['Age_calcule'].apply(
                lambda x: 'Proche retraite (55+)' if x >= 55 
                else 'Mi-carrière (35-54)' if x >= 35 
                else 'Jeune talent (<35)')
        
        if 'Anciennete_calculee' in df.columns:
            df['Segment_Anciennete'] = pd.cut(df['Anciennete_calculee'], 
                                            bins=[-1, 2, 5, 10, 20, 100], 
                                            labels=['Nouveau (0-2 ans)', 'Junior (2-5 ans)', 
                                                   'Expérimenté (5-10 ans)', 'Senior (10-20 ans)', 
                                                   'Expert (20+ ans)'])
        
        # Calcul des indicateurs de risque
        if 'Age_calcule' in df.columns and 'Anciennete_calculee' in df.columns:
            df['Risque_Depart'] = df.apply(lambda row: 
                'Élevé' if (row['Age_calcule'] >= 55 or row['Anciennete_calculee'] <= 1) 
                else 'Moyen' if (row['Age_calcule'] >= 45 and row['Anciennete_calculee'] <= 3)
                else 'Faible', axis=1)
        
        return df
    
    except Exception as e:
        st.error(f"Erreur lors du chargement des données: {e}")
        return None

# Fonction pour créer des métriques avancées
def create_advanced_metrics(df):
    """Calcule des métriques RH avancées"""
    metrics = {}
    
    if len(df) > 0:
        # Calculs de base
        metrics['total_employees'] = len(df)
        metrics['avg_age'] = df['Age_calcule'].mean() if 'Age_calcule' in df.columns else 0
        metrics['avg_tenure'] = df['Anciennete_calculee'].mean() if 'Anciennete_calculee' in df.columns else 0
        
        # Calculs avancés
        if 'Sexe' in df.columns:
            metrics['gender_ratio'] = (df['Sexe'] == 'Masculin').sum() / len(df) * 100
            metrics['diversity_index'] = 1 - ((df['Sexe'] == 'Masculin').sum() / len(df))**2 - ((df['Sexe'] == 'Féminin').sum() / len(df))**2
        
        if 'Age_calcule' in df.columns:
            metrics['retirement_risk'] = (df['Age_calcule'] >= 55).sum()
            metrics['young_talent'] = (df['Age_calcule'] <= 35).sum()
            metrics['median_age'] = df['Age_calcule'].median()
        
        if 'Anciennete_calculee' in df.columns:
            metrics['turnover_risk'] = (df['Anciennete_calculee'] <= 2).sum()
            metrics['experienced_staff'] = (df['Anciennete_calculee'] >= 10).sum()
        
        if 'Observation' in df.columns:
            departures = df[df['Observation'].notna() & (df['Observation'] != '')]
            metrics['turnover_rate'] = len(departures) / len(df) * 100 if len(df) > 0 else 0
        else:
            metrics['turnover_rate'] = 0
    
    return metrics

# Fonction pour créer des graphiques avancés
def create_advanced_visualizations(df):
    """Crée des visualisations avancées"""
    visualizations = {}
    
    # 1. Heatmap de répartition par âge et ancienneté
    if 'Age_calcule' in df.columns and 'Anciennete_calculee' in df.columns:
        age_bins = pd.cut(df['Age_calcule'], bins=5)
        tenure_bins = pd.cut(df['Anciennete_calculee'], bins=5)
        heatmap_data = pd.crosstab(age_bins, tenure_bins, normalize='index') * 100
        
        fig_heatmap = px.imshow(heatmap_data.values,
                               x=[f"{int(interval.left)}-{int(interval.right)}" for interval in heatmap_data.columns],
                               y=[f"{int(interval.left)}-{int(interval.right)}" for interval in heatmap_data.index],
                               color_continuous_scale='RdYlBu_r',
                               title="Matrice Âge vs Ancienneté (%)")
        fig_heatmap.update_layout(xaxis_title="Ancienneté (années)", yaxis_title="Âge (années)")
        visualizations['heatmap'] = fig_heatmap
    
    # 2. Graphique en radar des compétences par direction
    if 'Direction' in df.columns and 'CSP' in df.columns:
        direction_csp = pd.crosstab(df['Direction'], df['CSP'], normalize='index') * 100
        if len(direction_csp) > 0:
            top_directions = direction_csp.head(5)
            
            fig_radar = go.Figure()
            
            for direction in top_directions.index:
                fig_radar.add_trace(go.Scatterpolar(
                    r=top_directions.loc[direction].values,
                    theta=top_directions.columns,
                    fill='toself',
                    name=direction[:20]
                ))
            
            fig_radar.update_layout(
                polar=dict(
                    radialaxis=dict(visible=True, range=[0, 100])
                ),
                showlegend=True,
                title="Profil des Compétences par Direction"
            )
            visualizations['radar'] = fig_radar
    
    return visualizations

# Fonction pour générer une attestation de travail
def generate_work_certificate(employee_data, template_path="attestation de travail.docx", custom_reference=None, debug_mode=False):
    """
    Génère une attestation de travail personnalisée à partir du template existant
    
    Args:
        employee_data: Données de l'employé
        template_path: Chemin vers le template Word
        custom_reference: Référence personnalisée (optionnel)
        debug_mode: Mode debug pour afficher les remplacements (optionnel)
    """
    if not DOCX_AVAILABLE:
        return None, "La bibliothèque python-docx n'est pas installée. Veuillez l'installer avec: pip install python-docx", []
    
    try:
        # Vérifier si le template existe
        if not os.path.exists(template_path):
            return None, f"Le template '{template_path}' n'existe pas. Veuillez vous assurer que le fichier est présent dans le dossier.", []
        
        # Charger le template existant
        doc = Document(template_path)
        
        # Préparer les données de remplacement basées sur votre template PROMASIDOR
        nom_complet = f"{employee_data.get('Nom', '')} {employee_data.get('Prenoms', '')}".strip()
        
        # Formatage des dates en français
        def format_date_french(date_obj):
            """Formate une date en français"""
            if pd.isna(date_obj):
                return "Non renseigné"
            try:
                mois_fr = {
                    1: 'janvier', 2: 'février', 3: 'mars', 4: 'avril', 5: 'mai', 6: 'juin',
                    7: 'juillet', 8: 'août', 9: 'septembre', 10: 'octobre', 11: 'novembre', 12: 'décembre'
                }
                return f"{date_obj.day} {mois_fr[date_obj.month]} {date_obj.year}"
            except:
                return str(date_obj)
        
        date_naissance = format_date_french(employee_data.get('Date de naissance'))
        date_entree = format_date_french(employee_data.get('DateEntree'))
        
        # Date du jour en français
        mois_fr = {
            1: 'Janvier', 2: 'Février', 3: 'Mars', 4: 'Avril', 5: 'Mai', 6: 'Juin',
            7: 'Juillet', 8: 'Août', 9: 'Septembre', 10: 'Octobre', 11: 'Novembre', 12: 'Décembre'
        }
        
        aujourd_hui = datetime.now()
        date_generation = f"{aujourd_hui.day} {mois_fr[aujourd_hui.month]} {aujourd_hui.year}"
        
        # Préparer la référence à utiliser
        if custom_reference:
            reference_finale = custom_reference
        else:
            reference_finale = f"{employee_data.get('Matricule', '1261')} (ADM/DRH/{aujourd_hui.year})"
        
        # Dictionnaire de remplacement pour les placeholders du template PROMASIDOR
        replacements = {
            # Informations spécifiques à votre template
            'CHOUIKRAT Smail': nom_complet,
            '27 Juin 1990': date_naissance,
            'Hussin dey ,Alger': employee_data.get('Lieu de naissance', employee_data.get('Adresse', 'Non renseigné')),
            '04 Septembre 2022': date_entree,
            'de Head of strategy Sales': f"de {employee_data.get('Poste', 'Non renseigné')}",
            'Head of strategy Sales': employee_data.get('Poste', 'Non renseigné'),
            
            # Placeholders génériques si ils existent
            '[NOM_COMPLET]': nom_complet,
            '[DATE_NAISSANCE]': date_naissance,
            '[LIEU_NAISSANCE]': employee_data.get('Lieu de naissance', employee_data.get('Adresse', 'Non renseigné')),
            '[POSTE]': employee_data.get('Poste', 'Non renseigné'),
            '[DATE_ENTREE]': date_entree,
            '[DATE_GENERATION]': date_generation,
            '[REFERENCE]': reference_finale,
            
            # Mise à jour de la date dans l'en-tête
            '23/06/2025': aujourd_hui.strftime('%d/%m/%Y'),
            'le 23/06/2025': f"le {aujourd_hui.strftime('%d/%m/%Y')}",
            
            # Mise à jour de la référence - plusieurs formats possibles
            '1261 (ADM/DRH/2025)': reference_finale,
            '1261 (ADM/DRH/2025': reference_finale,  # Sans parenthèse fermante
            'Réf : 1261 (ADM/DRH/2025)': f"Réf : {reference_finale}",
            'Ref : 1261 (ADM/DRH/2025)': f"Ref : {reference_finale}",
            'N° : 1261 (ADM/DRH/2025)': f"N° : {reference_finale}",
            'Référence : 1261 (ADM/DRH/2025)': f"Référence : {reference_finale}",
        }
        
        # Fonction pour remplacer le texte et mettre en gras TOUTES les valeurs remplacées
        def replace_text_in_paragraph(paragraph, old_text, new_text, force_bold=True):
            if old_text in paragraph.text:
                full_text = paragraph.text
                if old_text in full_text:
                    # Sauvegarder le formatage original
                    original_formatting = {}
                    if paragraph.runs:
                        original_formatting = {
                            'font_name': paragraph.runs[0].font.name,
                            'font_size': paragraph.runs[0].font.size,
                            'bold': paragraph.runs[0].bold,
                            'italic': paragraph.runs[0].italic
                        }
                    
                    # Diviser le texte en parties
                    parts = full_text.split(old_text)
                    
                    # Nettoyer tous les runs
                    for run in paragraph.runs:
                        run.clear()
                    
                    # Reconstruire le paragraphe
                    for i, part in enumerate(parts):
                        if part:  # Ajouter la partie normale avec formatage original
                            normal_run = paragraph.add_run(part)
                            if original_formatting.get('font_name'):
                                normal_run.font.name = original_formatting['font_name']
                            if original_formatting.get('font_size'):
                                normal_run.font.size = original_formatting['font_size']
                            normal_run.bold = original_formatting.get('bold', False)
                            normal_run.italic = original_formatting.get('italic', False)
                        
                        if i < len(parts) - 1:  # Ajouter le texte de remplacement EN GRAS
                            replacement_run = paragraph.add_run(new_text)
                            if original_formatting.get('font_name'):
                                replacement_run.font.name = original_formatting['font_name']
                            if original_formatting.get('font_size'):
                                replacement_run.font.size = original_formatting['font_size']
                            # TOUTES les valeurs remplacées sont maintenant en gras
                            replacement_run.bold = True
                            replacement_run.italic = original_formatting.get('italic', False)
                    
                    return True
            return False
        
        # Fonction spéciale pour les cellules de tableau - TOUTES les valeurs en gras
        def replace_text_in_cell(cell, old_text, new_text, force_bold=True):
            for paragraph in cell.paragraphs:
                if old_text in paragraph.text:
                    full_text = paragraph.text
                    if old_text in full_text:
                        # Sauvegarder le formatage original
                        original_formatting = {}
                        if paragraph.runs:
                            original_formatting = {
                                'font_name': paragraph.runs[0].font.name,
                                'font_size': paragraph.runs[0].font.size,
                                'bold': paragraph.runs[0].bold,
                                'italic': paragraph.runs[0].italic
                            }
                        
                        # Diviser et reconstruire
                        parts = full_text.split(old_text)
                        
                        # Nettoyer
                        for run in paragraph.runs:
                            run.clear()
                        
                        # Reconstruire
                        for i, part in enumerate(parts):
                            if part:  # Texte normal avec formatage original
                                normal_run = paragraph.add_run(part)
                                if original_formatting.get('font_name'):
                                    normal_run.font.name = original_formatting['font_name']
                                if original_formatting.get('font_size'):
                                    normal_run.font.size = original_formatting['font_size']
                                normal_run.bold = original_formatting.get('bold', False)
                                normal_run.italic = original_formatting.get('italic', False)
                            
                            if i < len(parts) - 1:  # Texte de remplacement EN GRAS
                                replacement_run = paragraph.add_run(new_text)
                                if original_formatting.get('font_name'):
                                    replacement_run.font.name = original_formatting['font_name']
                                if original_formatting.get('font_size'):
                                    replacement_run.font.size = original_formatting['font_size']
                                # TOUTES les valeurs remplacées sont maintenant en gras
                                replacement_run.bold = True
                                replacement_run.italic = original_formatting.get('italic', False)
                        return True
            return False
        
        # Variables pour le debug
        replacements_made = []
        
        # Remplacer dans tous les paragraphes - TOUTES les valeurs en gras
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if replace_text_in_paragraph(paragraph, old_text, new_text):
                    if debug_mode:
                        replacements_made.append(f"Paragraphe: '{old_text}' → '{new_text}'")
        
        # Remplacer dans les tableaux si présents - TOUTES les valeurs en gras
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if replace_text_in_cell(cell, old_text, new_text):
                            if debug_mode:
                                replacements_made.append(f"Tableau: '{old_text}' → '{new_text}'")
        
        # Remplacer dans les headers et footers - TOUTES les valeurs en gras
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for old_text, new_text in replacements.items():
                    if replace_text_in_paragraph(paragraph, old_text, new_text):
                        if debug_mode:
                            replacements_made.append(f"En-tête: '{old_text}' → '{new_text}'")
            
            # Footer
            for paragraph in section.footer.paragraphs:
                for old_text, new_text in replacements.items():
                    if replace_text_in_paragraph(paragraph, old_text, new_text):
                        if debug_mode:
                            replacements_made.append(f"Pied de page: '{old_text}' → '{new_text}'")
        
        # Sauvegarder dans un buffer mémoire
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Retourner avec les informations de debug
        if debug_mode:
            return doc_buffer, None, replacements_made
        else:
            return doc_buffer, None, []
        
    except Exception as e:
        if debug_mode:
            return None, f"Erreur lors de la génération de l'attestation : {str(e)}", []
        else:
            return None, f"Erreur lors de la génération de l'attestation : {str(e)}", []

# Fonction pour créer la pyramide des âges
def create_age_pyramid(df):
    """Crée une pyramide des âges par sexe"""
    if 'Age_calcule' in df.columns and 'Sexe' in df.columns:
        # Créer des tranches d'âge
        bins = range(20, 70, 5)
        df['Tranche_age'] = pd.cut(df['Age_calcule'], bins=bins, right=False)
        
        # Compter par sexe et tranche d'âge
        pyramid_data = df.groupby(['Tranche_age', 'Sexe']).size().unstack(fill_value=0)
        
        if 'Masculin' in pyramid_data.columns:
            pyramid_data['Masculin'] = -pyramid_data['Masculin']  # Valeurs négatives pour les hommes
        
        fig = go.Figure()
        
        # Ajouter les barres pour les femmes (droite)
        if 'Féminin' in pyramid_data.columns:
            fig.add_trace(go.Bar(
                y=[str(interval) for interval in pyramid_data.index],
                x=pyramid_data['Féminin'],
                orientation='h',
                name='Féminin',
                marker_color='pink',
                text=pyramid_data['Féminin'],
                textposition='outside'
            ))
        
        # Ajouter les barres pour les hommes (gauche)
        if 'Masculin' in pyramid_data.columns:
            fig.add_trace(go.Bar(
                y=[str(interval) for interval in pyramid_data.index],
                x=pyramid_data['Masculin'],
                orientation='h',
                name='Masculin',
                marker_color='lightblue',
                text=abs(pyramid_data['Masculin']),
                textposition='outside'
            ))
        
        fig.update_layout(
            title='Pyramide des âges par sexe',
            xaxis_title='Nombre d\'employés',
            yaxis_title='Tranches d\'âge',
            barmode='relative',
            height=600
        )
        
        return fig
    return None

# Fonction principale de l'application
def main():
    # Header avec logo et titre
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        <div style='text-align: center; padding: 20px;'>
            <h1 style='color: #2c3e50; font-size: 2.5em; margin-bottom: 10px;'>
                DASHBOARD RH EXECUTIVE
            </h1>
            <p style='color: #7f8c8d; font-size: 1.2em; font-style: italic;'>
                Analytics & Business Intelligence Platform
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Chargement des données avec indicateur de progression
    with st.spinner('Chargement et analyse des données RH...'):
        df = load_and_clean_data()
    
    if df is None:
        st.error("Impossible de charger les données. Vérifiez le fichier source.")
        st.stop()
    
    # Message de succès avec informations sur les données
    st.markdown(f"""
    <div class="success-box">
        <h4>Données chargées avec succès</h4>
        <p><strong>Total des enregistrements:</strong> {len(df)} employés</p>
        <p><strong>Dernière mise à jour:</strong> {datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
        <p><strong>Période couverte:</strong> {df['DateEntree'].min().strftime('%Y') if 'DateEntree' in df.columns else 'N/A'} - {datetime.now().year}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar améliorée avec design professionnel
    st.sidebar.markdown("""
    <div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                padding: 20px; border-radius: 10px; margin-bottom: 20px;'>
        <h2 style='color: white; text-align: center; margin-bottom: 15px;'>
            FILTRES & CONTRÔLES
        </h2>
        <p style='color: #f8f9fa; text-align: center; font-size: 0.9em;'>
            Personnalisez votre analyse
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Filtres avancés
    st.sidebar.subheader("Filtres Organisationnels")
    
    directions = ['Toutes les directions'] + [f"{dir}" for dir in sorted(df['Direction'].dropna().unique().tolist())]
    selected_direction = st.sidebar.selectbox("Direction", directions)
    
    sexes = ['Tous les sexes'] + [f"{sexe}" for sexe in sorted(df['Sexe'].dropna().unique().tolist())]
    selected_sexe = st.sidebar.selectbox("Sexe", sexes)
    
    types_contrat = ['Tous les contrats'] + [f"{contrat}" for contrat in sorted(df['Type de contrat'].dropna().unique().tolist())]
    selected_contrat = st.sidebar.selectbox("Type de contrat", types_contrat)
    
    # Filtres par tranche d'âge
    if 'Age_calcule' in df.columns:
        st.sidebar.subheader("Filtres Démographiques")
        age_range = st.sidebar.slider(
            "Tranche d'âge",
            min_value=int(df['Age_calcule'].min()),
            max_value=int(df['Age_calcule'].max()),
            value=(int(df['Age_calcule'].min()), int(df['Age_calcule'].max())),
            step=1
        )
    
    # Filtres par ancienneté
    if 'Anciennete_calculee' in df.columns:
        tenure_range = st.sidebar.slider(
            "Ancienneté (années)",
            min_value=0.0,
            max_value=float(df['Anciennete_calculee'].max()),
            value=(0.0, float(df['Anciennete_calculee'].max())),
            step=0.5
        )
    
    # Application des filtres
    filtered_df = df.copy()
    
    if not selected_direction.startswith('Toutes'):
        filtered_df = filtered_df[filtered_df['Direction'] == selected_direction]
    
    if not selected_sexe.startswith('Tous'):
        filtered_df = filtered_df[filtered_df['Sexe'] == selected_sexe]
    
    if not selected_contrat.startswith('Tous'):
        filtered_df = filtered_df[filtered_df['Type de contrat'] == selected_contrat]
    
    if 'Age_calcule' in filtered_df.columns:
        filtered_df = filtered_df[
            (filtered_df['Age_calcule'] >= age_range[0]) & 
            (filtered_df['Age_calcule'] <= age_range[1])
        ]
    
    if 'Anciennete_calculee' in filtered_df.columns:
        filtered_df = filtered_df[
            (filtered_df['Anciennete_calculee'] >= tenure_range[0]) & 
            (filtered_df['Anciennete_calculee'] <= tenure_range[1])
        ]
    
    # Affichage des filtres actifs
    active_filters = []
    if not selected_direction.startswith('Toutes'):
        active_filters.append(f"Direction: {selected_direction}")
    if not selected_sexe.startswith('Tous'):
        active_filters.append(f"Sexe: {selected_sexe}")
    if not selected_contrat.startswith('Tous'):
        active_filters.append(f"Contrat: {selected_contrat}")
    
    if active_filters:
        st.sidebar.markdown(f"""
        <div class="info-box">
            <h5>Filtres Actifs:</h5>
            {'<br>'.join([f"• {filter}" for filter in active_filters])}
        </div>
        """, unsafe_allow_html=True)
    
    # Section d'information sur l'attestation de travail
    if not DOCX_AVAILABLE:
        st.sidebar.markdown("""
        <div style='background: linear-gradient(135deg, #ff6b6b 0%, #ffa500 100%); 
                    padding: 15px; border-radius: 10px; margin-top: 20px;'>
            <h4 style='color: white; margin-bottom: 10px;'>🔧 Installation Requise</h4>
            <p style='color: #f8f9fa; font-size: 0.9em; margin-bottom: 10px;'>
                Pour utiliser le générateur d'attestation, exécutez :
            </p>
            <p style='color: #f8f9fa; font-size: 0.8em; background: rgba(0,0,0,0.2); 
                      padding: 8px; border-radius: 5px; margin: 0;'>
                install_attestation.bat
            </p>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.sidebar.markdown("""
        <div style='background: linear-gradient(135deg, #28a745 0%, #20c997 100%); 
                    padding: 15px; border-radius: 10px; margin-top: 20px;'>
            <h4 style='color: white; margin-bottom: 10px;'>📋 Attestations PROMASIDOR</h4>
            <p style='color: #f8f9fa; font-size: 0.9em; margin: 0;'>
                Générateur d'attestation avec votre template officiel !
            </p>
            <p style='color: #e9ecef; font-size: 0.8em; margin-top: 5px; margin-bottom: 0;'>
                Utilise "attestation de travail.docx" avec logo et formatage PROMASIDOR.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Calcul des métriques avancées
    metrics = create_advanced_metrics(filtered_df)
    
    # Dashboard principal avec métriques en cards
    st.header("TABLEAU DE BORD EXÉCUTIF")
    
    # Première ligne de métriques principales
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        st.markdown("""
        <div class="metric-container" style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);">
            <h3 style="margin: 0; color: white;">Employés</h3>
            <h2 style="margin: 5px 0; color: white;">{}</h2>
            <p style="margin: 0; color: #f8f9fa;">Total</p>
        </div>
        """.format(metrics['total_employees']), unsafe_allow_html=True)
    
    with col2:
        avg_age = metrics['avg_age']
        age_trend = "ÉLEVÉ" if avg_age > 45 else "FAIBLE" if avg_age < 35 else "MOYEN"
        st.markdown("""
        <div class="metric-container" style="background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);">
            <h3 style="margin: 0; color: #8b4513;">{}</h3>
            <h2 style="margin: 5px 0; color: #8b4513;">{:.1f} ans</h2>
            <p style="margin: 0; color: #a0522d;">Âge Moyen</p>
        </div>
        """.format(age_trend, avg_age), unsafe_allow_html=True)
    
    with col3:
        avg_tenure = metrics['avg_tenure']
        tenure_icon = "EXPERT" if avg_tenure > 10 else "CONFIRMÉ" if avg_tenure > 5 else "JUNIOR"
        st.markdown("""
        <div class="metric-container" style="background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);">
            <h3 style="margin: 0; color: #2c3e50;">{}</h3>
            <h2 style="margin: 5px 0; color: #2c3e50;">{:.1f} ans</h2>
            <p style="margin: 0; color: #34495e;">Ancienneté Moy.</p>
        </div>
        """.format(tenure_icon, avg_tenure), unsafe_allow_html=True)
    
    with col4:
        gender_ratio = metrics['gender_ratio']
        diversity_level = "FAIBLE" if gender_ratio > 80 or gender_ratio < 20 else "MOYEN" if gender_ratio > 70 or gender_ratio < 30 else "BON"
        st.markdown("""
        <div class="metric-container" style="background: linear-gradient(135deg, #ff9a9e 0%, #fecfef 100%);">
            <h3 style="margin: 0; color: #8b0000;">{}</h3>
            <h2 style="margin: 5px 0; color: #8b0000;">{:.1f}%</h2>
            <p style="margin: 0; color: #a0522d;">Taux Masculin</p>
        </div>
        """.format(diversity_level, gender_ratio), unsafe_allow_html=True)
    
    with col5:
        turnover = metrics['turnover_rate']
        turnover_status = "ÉLEVÉ" if turnover > 15 else "MOYEN" if turnover > 10 else "FAIBLE"
        st.markdown("""
        <div class="metric-container" style="background: linear-gradient(135deg, #fa709a 0%, #fee140 100%);">
            <h3 style="margin: 0; color: #8b0000;">{}</h3>
            <h2 style="margin: 5px 0; color: #8b0000;">{:.1f}%</h2>
            <p style="margin: 0; color: #a0522d;">Taux Rotation</p>
        </div>
        """.format(turnover_status, turnover), unsafe_allow_html=True)
    
    # Deuxième ligne de métriques de risque
    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        retirement_risk = metrics.get('retirement_risk', 0)
        st.metric(
            "Risque Retraite",
            f"{retirement_risk} employés",
            f"{retirement_risk/len(filtered_df)*100:.1f}% des effectifs" if len(filtered_df) > 0 else "0%"
        )
    
    with col2:
        young_talent = metrics.get('young_talent', 0)
        st.metric(
            "Jeunes Talents",
            f"{young_talent} employés",
            f"{young_talent/len(filtered_df)*100:.1f}% des effectifs" if len(filtered_df) > 0 else "0%"
        )
    
    with col3:
        experienced = metrics.get('experienced_staff', 0)
        st.metric(
            "Personnel Expérimenté",
            f"{experienced} employés",
            f"{experienced/len(filtered_df)*100:.1f}% des effectifs" if len(filtered_df) > 0 else "0%"
        )
    
    with col4:
        diversity_index = metrics.get('diversity_index', 0)
        diversity_score = "Excellent" if diversity_index > 0.4 else "Bon" if diversity_index > 0.3 else "À améliorer"
        st.metric(
            "Index Diversité",
            f"{diversity_score}",
            f"Score: {diversity_index:.2f}"
        )
    
    st.markdown("---")
    
    
    # Section des visualisations principales
    st.header("ANALYSES VISUELLES AVANCÉES")
    
    # Onglets pour organiser les analyses
    tab1, tab2, tab3, tab4 = st.tabs(["Démographie", "Organisation", "Performance", "Analyses Avancées"])
    
    with tab1:
        st.subheader("Analyse Démographique Complète")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Graphique en secteurs amélioré pour le sexe
            if 'Sexe' in filtered_df.columns:
                sexe_counts = filtered_df['Sexe'].value_counts()
                fig_sexe = px.pie(
                    values=sexe_counts.values, 
                    names=sexe_counts.index,
                    title="Répartition par Genre",
                    color_discrete_map={'Masculin': '#3498db', 'Féminin': '#e74c3c'},
                    hole=0.4
                )
                fig_sexe.update_traces(textposition='inside', textinfo='percent+label')
                fig_sexe.update_layout(
                    showlegend=True,
                    font=dict(size=12),
                    title_font_size=16
                )
                st.plotly_chart(fig_sexe, use_container_width=True, key="sexe_pie_chart")
                
                # Analyse de la diversité
                diversity_ratio = min(sexe_counts.values) / max(sexe_counts.values) * 100
                if diversity_ratio > 40:
                    diversity_status = "Excellente diversité"
                elif diversity_ratio > 25:
                    diversity_status = "Diversité correcte"
                else:
                    diversity_status = "Diversité à améliorer"
                
                st.markdown(f"""
                <div class="info-box">
                    <strong>Analyse de Diversité:</strong><br>
                    {diversity_status}<br>
                    <em>Ratio de diversité: {diversity_ratio:.1f}%</em>
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            # Pyramide des âges modernisée
            if 'Age_calcule' in filtered_df.columns and 'Sexe' in filtered_df.columns:
                # Créer des tranches d'âge
                bins = range(20, 70, 5)
                filtered_df_temp = filtered_df.copy()
                filtered_df_temp['Tranche_age'] = pd.cut(filtered_df_temp['Age_calcule'], bins=bins, right=False)
                
                # Compter par sexe et tranche d'âge
                pyramid_data = filtered_df_temp.groupby(['Tranche_age', 'Sexe']).size().unstack(fill_value=0)
                
                if len(pyramid_data) > 0:
                    fig_pyramid = go.Figure()
                    
                    # Hommes (à gauche, valeurs négatives)
                    if 'Masculin' in pyramid_data.columns:
                        fig_pyramid.add_trace(go.Bar(
                            y=[str(interval) for interval in pyramid_data.index],
                            x=-pyramid_data['Masculin'],
                            orientation='h',
                            name='Masculin',
                            marker_color='#3498db',
                            text=pyramid_data['Masculin'],
                            textposition='outside',
                            hovertemplate='Masculin: %{text}<extra></extra>'
                        ))
                    
                    # Femmes (à droite, valeurs positives)
                    if 'Féminin' in pyramid_data.columns:
                        fig_pyramid.add_trace(go.Bar(
                            y=[str(interval) for interval in pyramid_data.index],
                            x=pyramid_data['Féminin'],
                            orientation='h',
                            name='Féminin',
                            marker_color='#e74c3c',
                            text=pyramid_data['Féminin'],
                            textposition='outside',
                            hovertemplate='Féminin: %{text}<extra></extra>'
                        ))
                    
                    fig_pyramid.update_layout(
                        title='Pyramide des Âges Interactive',
                        xaxis_title='Nombre d\'employés',
                        yaxis_title='Tranches d\'âge',
                        barmode='relative',
                        height=500,
                        showlegend=True
                    )
                    
                    st.plotly_chart(fig_pyramid, use_container_width=True, key="age_pyramid_chart")
        
        # Analyse des générations
        col1, col2 = st.columns(2)
        
        with col1:
            if 'Generation' in filtered_df.columns:
                gen_counts = filtered_df['Generation'].value_counts()
                fig_gen = px.bar(
                    x=gen_counts.index,
                    y=gen_counts.values,
                    title="Répartition par Génération",
                    color=gen_counts.values,
                    color_continuous_scale='viridis'
                )
                fig_gen.update_layout(
                    xaxis_title="Génération",
                    yaxis_title="Nombre d'employés",
                    showlegend=False
                )
                st.plotly_chart(fig_gen, use_container_width=True, key="generation_bar_chart")
        
        with col2:
            if 'Statut_Retraite' in filtered_df.columns:
                retraite_counts = filtered_df['Statut_Retraite'].value_counts()
                fig_retraite = px.pie(
                    values=retraite_counts.values,
                    names=retraite_counts.index,
                    title="Statut de Carrière",
                    color_discrete_sequence=['#2ecc71', '#f39c12', '#e74c3c']
                )
                st.plotly_chart(fig_retraite, use_container_width=True, key="statut_carriere_pie_chart")
    
    with tab2:
        st.subheader("Analyse Organisationnelle")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Top directions avec design amélioré
            if 'Direction' in filtered_df.columns:
                direction_counts = filtered_df['Direction'].value_counts().head(10)
                fig_direction = px.bar(
                    y=direction_counts.index,
                    x=direction_counts.values,
                    orientation='h',
                    title="Top 10 des Directions",
                    color=direction_counts.values,
                    color_continuous_scale='blues'
                )
                fig_direction.update_layout(
                    xaxis_title="Nombre d'employés",
                    yaxis_title="Direction",
                    height=500
                )
                st.plotly_chart(fig_direction, use_container_width=True, key="directions_bar_chart")
        
        with col2:
            # Types de contrat avec indicateurs
            if 'Type de contrat' in filtered_df.columns:
                contrat_counts = filtered_df['Type de contrat'].value_counts()
                fig_contrat = px.pie(
                    values=contrat_counts.values,
                    names=contrat_counts.index,
                    title="Types de Contrat",
                    hole=0.3,
                    color_discrete_sequence=px.colors.qualitative.Set3
                )
                st.plotly_chart(fig_contrat, use_container_width=True, key="contrats_pie_chart")
                
                # Analyse de stabilité contractuelle
                cdi_ratio = (filtered_df['Type de contrat'] == 'CDI').sum() / len(filtered_df) * 100 if len(filtered_df) > 0 else 0
                if cdi_ratio > 90:
                    stability_status = "Stabilité Excellente"
                elif cdi_ratio > 70:
                    stability_status = "Stabilité Correcte"
                else:
                    stability_status = "Stabilité Faible"
                
                st.markdown(f"""
                <div class="success-box">
                    <strong>Stabilité Contractuelle:</strong><br>
                    {stability_status}<br>
                    <em>Taux CDI: {cdi_ratio:.1f}%</em>
                </div>
                """, unsafe_allow_html=True)
        
        # CSP et situation civile
        col1, col2 = st.columns(2)
        
        with col1:
            if 'CSP' in filtered_df.columns:
                csp_counts = filtered_df['CSP'].value_counts()
                fig_csp = px.bar(
                    x=csp_counts.index,
                    y=csp_counts.values,
                    title="Catégories Socio-Professionnelles",
                    color=csp_counts.values,
                    color_continuous_scale='plasma'
                )
                st.plotly_chart(fig_csp, use_container_width=True, key="csp_bar_chart")
        
        with col2:
            if 'Situation Civile' in filtered_df.columns:
                situation_counts = filtered_df['Situation Civile'].value_counts()
                fig_situation = px.pie(
                    values=situation_counts.values,
                    names=situation_counts.index,
                    title="Situation Civile",
                    hole=0.3,
                    color_discrete_sequence=px.colors.qualitative.Pastel
                )
                st.plotly_chart(fig_situation, use_container_width=True, key="situation_civile_pie_chart")
    
    with tab3:
        st.subheader("Analyse de Performance et Risques")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Distribution des âges avec statistiques
            if 'Age_calcule' in filtered_df.columns:
                fig_age = px.histogram(
                    filtered_df,
                    x='Age_calcule',
                    nbins=25,
                    title="Distribution des Âges",
                    color_discrete_sequence=['#3498db']
                )
                
                # Ajouter les lignes de moyenne et médiane
                mean_age = filtered_df['Age_calcule'].mean()
                median_age = filtered_df['Age_calcule'].median()
                
                fig_age.add_vline(x=mean_age, line_dash="dash", line_color="red", 
                                annotation_text=f"Moyenne: {mean_age:.1f}")
                fig_age.add_vline(x=median_age, line_dash="dot", line_color="green",
                                annotation_text=f"Médiane: {median_age:.1f}")
                
                fig_age.update_layout(
                    xaxis_title="Âge (années)",
                    yaxis_title="Nombre d'employés"
                )
                st.plotly_chart(fig_age, use_container_width=True, key="age_histogram_chart")
        
        with col2:
            # Analyse des risques de départ
            if 'Risque_Depart' in filtered_df.columns:
                risk_counts = filtered_df['Risque_Depart'].value_counts()
                colors = {'Faible': '#2ecc71', 'Moyen': '#f39c12', 'Élevé': '#e74c3c'}
                fig_risk = px.pie(
                    values=risk_counts.values,
                    names=risk_counts.index,
                    title="Analyse des Risques de Départ",
                    color=risk_counts.index,
                    color_discrete_map=colors
                )
                st.plotly_chart(fig_risk, use_container_width=True, key="risque_depart_pie_chart")
        
        # Analyse des segments d'ancienneté
        if 'Segment_Anciennete' in filtered_df.columns:
            anc_counts = filtered_df['Segment_Anciennete'].value_counts()
            fig_anc = px.bar(
                x=anc_counts.index,
                y=anc_counts.values,
                title="Répartition par Ancienneté",
                color=anc_counts.values,
                color_continuous_scale='YlOrRd'
            )
            fig_anc.update_layout(
                xaxis_title="Segment d'ancienneté",
                yaxis_title="Nombre d'employés"
            )
            st.plotly_chart(fig_anc, use_container_width=True, key="anciennete_bar_chart")
    
    with tab4:
        st.subheader("Analyses Avancées et Insights")
        
        # Créer les visualisations avancées
        advanced_viz = create_advanced_visualizations(filtered_df)
        
        # Heatmap âge vs ancienneté
        if 'heatmap' in advanced_viz:
            st.plotly_chart(advanced_viz['heatmap'], use_container_width=True, key="heatmap_age_anciennete")
        
        # Graphique radar des compétences
        if 'radar' in advanced_viz:
            st.plotly_chart(advanced_viz['radar'], use_container_width=True, key="radar_competences")
        
        # Analyse prédictive des départs en retraite
        if 'Age_calcule' in filtered_df.columns:
            st.subheader("Prévisions de Départs en Retraite")
            
            current_year = datetime.now().year
            retirement_forecast = []
            
            for year_offset in range(1, 6):
                future_year = current_year + year_offset
                future_age = filtered_df['Age_calcule'] + year_offset
                retirees = (future_age >= 62).sum()
                retirement_forecast.append({'Année': future_year, 'Départs Prévus': retirees})
            
            forecast_df = pd.DataFrame(retirement_forecast)
            
            fig_forecast = px.line(
                forecast_df,
                x='Année',
                y='Départs Prévus',
                title="Prévisions de Départs en Retraite (5 ans)",
                markers=True
            )
            fig_forecast.update_traces(line_color='#e74c3c', line_width=3)
            st.plotly_chart(fig_forecast, use_container_width=True, key="forecast_retraite_chart")
            
            # Table des prévisions
            st.dataframe(forecast_df, use_container_width=True)
    
    
    # Analyse des départs améliorée
    st.header("ANALYSE AVANCÉE DES DÉPARTS")
    if 'Observation' in filtered_df.columns:
        departs = filtered_df[filtered_df['Observation'].notna() & (filtered_df['Observation'] != '')]
        if len(departs) > 0:
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown(f"""
                <div class="highlight-box">
                    <h4>Statistiques de Rotation</h4>
                    <ul>
                        <li><strong>Total départs:</strong> {len(departs)}</li>
                        <li><strong>Taux de rotation:</strong> {len(departs)/len(df)*100:.1f}%</li>
                        <li><strong>Impact sur effectif:</strong> {len(departs)/len(filtered_df)*100:.1f}%</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                # Analyse des raisons de départ
                raisons_depart = departs['Observation'].value_counts()
                fig_departs = px.pie(
                    values=raisons_depart.values,
                    names=raisons_depart.index,
                    title="Raisons de Départ",
                    color_discrete_sequence=px.colors.qualitative.Set3
                )
                st.plotly_chart(fig_departs, use_container_width=True, key="raisons_depart_pie_chart")
            
            with col3:
                # Profil des partants par âge
                if 'Age_calcule' in departs.columns:
                    fig_age_departs = px.histogram(
                        departs,
                        x='Age_calcule',
                        title="Âge des Partants",
                        nbins=15,
                        color_discrete_sequence=['#e74c3c']
                    )
                    st.plotly_chart(fig_age_departs, use_container_width=True, key="age_partants_histogram")
            
            # Tableau détaillé des départs
            st.subheader("Détail des Départs")
            depart_columns = ['Nom', 'Prenoms', 'Age_calcule', 'Poste', 'Direction', 'DateEntree', 'Anciennete_calculee', 'Observation']
            available_depart_cols = [col for col in depart_columns if col in departs.columns]
            
            if available_depart_cols:
                departs_display = departs[available_depart_cols].copy()
                departs_display = departs_display.sort_values('Anciennete_calculee', ascending=False)
                st.dataframe(departs_display, use_container_width=True)
        else:
            st.markdown("""
            <div class="success-box">
                <h4>Excellent Taux de Rétention</h4>
                <p>Aucun départ enregistré dans les données filtrées. 
                Cela indique une excellente stabilité de l'équipe et un environnement de travail favorable.</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Section d'analytics basées sur les données réelles
    st.header("INDICATEURS CLÉS")
    
    # Métriques simples et claires
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        # Répartition par sexe
        if 'Sexe' in filtered_df.columns:
            masculin_count = (filtered_df['Sexe'] == 'Masculin').sum()
            feminin_count = (filtered_df['Sexe'] == 'Féminin').sum()
            
            st.markdown(f"""
            <div class="metric-container" style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);">
                <h4 style="margin: 0; color: white;">Répartition Genre</h4>
                <h2 style="margin: 5px 0; color: white;">{masculin_count}H / {feminin_count}F</h2>
                <p style="margin: 0; color: #f8f9fa;">Effectifs par sexe</p>
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        # Répartition par type de contrat
        if 'Type de contrat' in filtered_df.columns:
            cdi_count = (filtered_df['Type de contrat'] == 'CDI').sum()
            total_count = len(filtered_df)
            cdi_percentage = (cdi_count / total_count * 100) if total_count > 0 else 0
            
            st.markdown(f"""
            <div class="metric-container" style="background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);">
                <h4 style="margin: 0; color: #8b4513;">Contrats CDI</h4>
                <h2 style="margin: 5px 0; color: #8b4513;">{cdi_count} ({cdi_percentage:.1f}%)</h2>
                <p style="margin: 0; color: #a0522d;">Sur {total_count} employés</p>
            </div>
            """, unsafe_allow_html=True)
    
    with col3:
        # Employés par tranche d'âge
        if 'Age_calcule' in filtered_df.columns:
            jeunes = (filtered_df['Age_calcule'] <= 35).sum()
            seniors = (filtered_df['Age_calcule'] >= 55).sum()
            
            st.markdown(f"""
            <div class="metric-container" style="background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);">
                <h4 style="margin: 0; color: #2c3e50;">Tranches d'Âge</h4>
                <h2 style="margin: 5px 0; color: #2c3e50;">{jeunes} Jeunes / {seniors} Seniors</h2>
                <p style="margin: 0; color: #34495e;">≤35 ans / ≥55 ans</p>
            </div>
            """, unsafe_allow_html=True)
    
    with col4:
        # Employés par ancienneté
        if 'Anciennete_calculee' in filtered_df.columns:
            nouveaux = (filtered_df['Anciennete_calculee'] <= 2).sum()
            anciens = (filtered_df['Anciennete_calculee'] >= 10).sum()
            
            st.markdown(f"""
            <div class="metric-container" style="background: linear-gradient(135deg, #ff9a9e 0%, #fecfef 100%);">
                <h4 style="margin: 0; color: #8b0000;">Ancienneté</h4>
                <h2 style="margin: 5px 0; color: #8b0000;">{nouveaux} Nouveaux / {anciens} Anciens</h2>
                <p style="margin: 0; color: #a0522d;">≤2 ans / ≥10 ans</p>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Analytics détaillées
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Analyse des Risques Réels")
        
        # Analyse basée sur les données réelles uniquement
        if 'Age_calcule' in filtered_df.columns and 'Anciennete_calculee' in filtered_df.columns:
            # Analyse des employés proches de la retraite (données réelles)
            proche_retraite = (filtered_df['Age_calcule'] >= 60).sum()
            nouveaux_employes = (filtered_df['Anciennete_calculee'] <= 1).sum()
            
            # Répartition par tranche d'ancienneté
            if 'Segment_Anciennete' in filtered_df.columns:
                anciennete_counts = filtered_df['Segment_Anciennete'].value_counts()
                
                fig_anciennete = px.bar(
                    x=anciennete_counts.index,
                    y=anciennete_counts.values,
                    title="Répartition par Ancienneté",
                    color=anciennete_counts.values,
                    color_continuous_scale='viridis'
                )
                fig_anciennete.update_layout(
                    xaxis_title="Segment d'Ancienneté",
                    yaxis_title="Nombre d'Employés"
                )
                st.plotly_chart(fig_anciennete, use_container_width=True, key="anciennete_risque_chart")
            
            # Statistiques réelles
            st.markdown(f"""
            <div class="info-box">
                <strong>Indicateurs Basés sur Données Réelles:</strong><br>
                • <strong>Proches de la retraite (60+):</strong> {proche_retraite} employés<br>
                • <strong>Nouveaux employés (≤1 an):</strong> {nouveaux_employes} employés<br>
                • <strong>Ancienneté moyenne:</strong> {filtered_df['Anciennete_calculee'].mean():.1f} ans
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.subheader("Répartition par Poste")
        
        # Analyse des postes (données réelles)
        if 'Poste' in filtered_df.columns:
            poste_counts = filtered_df['Poste'].value_counts().head(10)
            
            fig_postes = px.bar(
                x=poste_counts.values,
                y=poste_counts.index,
                orientation='h',
                title="Top 10 des Postes",
                color=poste_counts.values,
                color_continuous_scale='plasma'
            )
            fig_postes.update_layout(
                xaxis_title="Nombre d'employés",
                yaxis_title="Poste",
                height=400
            )
            st.plotly_chart(fig_postes, use_container_width=True, key="postes_analytics_chart")
            
            # Statistiques des postes
            st.markdown(f"""
            <div class="info-box">
                <strong>Analyse des Postes:</strong><br>
                • <strong>Nombre de postes différents:</strong> {len(filtered_df['Poste'].unique())}<br>
                • <strong>Poste le plus courant:</strong> {poste_counts.index[0]}<br>
                • <strong>Employés dans ce poste:</strong> {poste_counts.iloc[0]} personnes
            </div>
            """, unsafe_allow_html=True)

    # SECTION GÉNÉRATION D'ATTESTATIONS DE TRAVAIL
    st.markdown("---")
    st.header("🏆 GÉNÉRATEUR D'ATTESTATIONS DE TRAVAIL")
    
    if DOCX_AVAILABLE:
        # Vérifier si le template existe
        template_path = "attestation de travail.docx"
        if os.path.exists(template_path):
            st.markdown("""
            <div style='background: linear-gradient(135deg, #28a745 0%, #20c997 100%); 
                        padding: 20px; border-radius: 10px; margin: 20px 0;'>
                <h3 style='color: white; margin-bottom: 15px;'>📋 Générateur PROMASIDOR</h3>
                <p style='color: #f8f9fa; margin: 0;'>
                    ✅ Template officiel détecté : <strong>attestation de travail.docx</strong><br>
                    ✅ Formatage automatique avec logo et en-tête PROMASIDOR<br>
                    ✅ Toutes les informations remplacées apparaîtront en <strong>GRAS</strong><br>
                    ✅ Dates formatées en français
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.subheader("Sélection de l'employé")
                
                # Créer une liste des employés avec nom complet
                employee_options = []
                employee_data_map = {}
                
                for index, row in filtered_df.iterrows():
                    nom_complet = f"{row.get('Nom', 'N/A')} {row.get('Prenoms', 'N/A')}".strip()
                    poste = row.get('Poste', 'Non renseigné')
                    direction = row.get('Direction', 'Non renseignée')
                    display_name = f"{nom_complet} - {poste} ({direction})"
                    employee_options.append(display_name)
                    employee_data_map[display_name] = row
                
                if employee_options:
                    selected_employee = st.selectbox(
                        "Choisir un employé pour générer son attestation :",
                        employee_options,
                        help="Sélectionnez l'employé pour lequel vous souhaitez générer une attestation de travail"
                    )
                    
                    if selected_employee:
                        employee_data = employee_data_map[selected_employee]
                        
                        # Afficher un aperçu des informations
                        st.markdown("### 👤 Aperçu des informations")
                        col_info1, col_info2 = st.columns(2)
                        
                        with col_info1:
                            st.markdown(f"""
                            <div class="info-box">
                                <strong>Informations personnelles :</strong><br>
                                • <strong>Nom complet :</strong> {employee_data.get('Nom', 'N/A')} {employee_data.get('Prenoms', 'N/A')}<br>
                                • <strong>Date de naissance :</strong> {employee_data.get('Date de naissance', 'N/A')}<br>
                                • <strong>Lieu de naissance :</strong> {employee_data.get('Lieu de naissance', employee_data.get('Adresse', 'Non renseigné'))}
                            </div>
                            """, unsafe_allow_html=True)
                        
                        with col_info2:
                            st.markdown(f"""
                            <div class="success-box">
                                <strong>Informations professionnelles :</strong><br>
                                • <strong>Poste :</strong> {employee_data.get('Poste', 'Non renseigné')}<br>
                                • <strong>Date d'entrée :</strong> {employee_data.get('DateEntree', 'N/A')}<br>
                                • <strong>Direction :</strong> {employee_data.get('Direction', 'Non renseignée')}
                            </div>
                            """, unsafe_allow_html=True)
            
            with col2:
                st.subheader("Paramètres de génération")
                
                # Option pour modifier la référence manuellement
                st.markdown("#### 📝 Référence du document")
                
                # Générer une référence par défaut
                default_reference = f"{employee_data.get('Matricule', '1261')} (ADM/DRH/{datetime.now().year})"
                
                # Checkbox pour personnaliser la référence
                customize_reference = st.checkbox(
                    "Personnaliser la référence",
                    help="Cochez pour modifier la référence par défaut"
                )
                
                if customize_reference:
                    custom_reference = st.text_input(
                        "Référence personnalisée :",
                        value=default_reference,
                        help="Saisissez la référence souhaitée pour ce document"
                    )
                    reference_to_use = custom_reference
                else:
                    reference_to_use = default_reference
                    st.info(f"📋 Référence par défaut : **{default_reference}**")
                
                # Option de debug
                debug_mode = st.checkbox(
                    "🔍 Mode debug (afficher les remplacements)",
                    help="Cochez pour voir quels remplacements sont effectués dans le document"
                )
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                if st.button("🎯 GÉNÉRER ATTESTATION", type="primary", use_container_width=True):
                    with st.spinner('Génération de l\'attestation en cours...'):
                        # Passer la référence personnalisée et le mode debug à la fonction
                        doc_buffer, error, debug_info = generate_work_certificate(
                            employee_data, 
                            template_path, 
                            custom_reference=reference_to_use,
                            debug_mode=debug_mode
                        )
                        
                        if doc_buffer and not error:
                            # Succès
                            st.markdown("""
                            <div class="success-box">
                                <h4>✅ Attestation générée avec succès !</h4>
                                <p>L'attestation a été créée avec les informations en <strong>GRAS</strong> et les dates en français.</p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Afficher les informations de debug si activées
                            if debug_mode and debug_info:
                                st.markdown("### 🔍 Détails des remplacements effectués")
                                if debug_info:
                                    for replacement in debug_info:
                                        st.write(f"✅ {replacement}")
                                else:
                                    st.warning("⚠️ Aucun remplacement détecté. Vérifiez que le template contient les placeholders attendus.")
                            
                            # Bouton de téléchargement
                            nom_fichier = f"Attestation_{employee_data.get('Nom', 'Employe')}_{employee_data.get('Prenoms', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
                            nom_fichier = nom_fichier.replace(' ', '_')
                            
                            st.download_button(
                                label="📥 Télécharger l'attestation",
                                data=doc_buffer.getvalue(),
                                file_name=nom_fichier,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="secondary",
                                use_container_width=True
                            )
                            
                            st.markdown(f"""
                            <div class="highlight-box">
                                <strong>📋 Informations sur le document :</strong><br>
                                • <strong>Nom du fichier :</strong> {nom_fichier}<br>
                                • <strong>Format :</strong> Microsoft Word (.docx)<br>
                                • <strong>Référence utilisée :</strong> {reference_to_use}<br>
                                • <strong>Formatage :</strong> Toutes les valeurs remplacées en GRAS<br>
                                • <strong>Dates :</strong> Format français (ex: "2 août 2025")
                            </div>
                            """, unsafe_allow_html=True)
                        
                        else:
                            # Erreur
                            st.error(f"❌ Erreur lors de la génération : {error}")
                
                # Informations sur le formatage
                st.markdown("### ℹ️ Formatage automatique")
                st.markdown(f"""
                <div style='background: #e3f2fd; padding: 15px; border-radius: 8px; border-left: 4px solid #2196f3;'>
                    <h5 style='color: #1976d2; margin-top: 0;'>Éléments personnalisés :</h5>
                    <ul style='color: #424242; margin-bottom: 10px;'>
                        <li>📝 <strong>Référence :</strong> {reference_to_use}</li>
                        <li>📅 <strong>Date :</strong> {datetime.now().strftime('%d/%m/%Y')}</li>
                    </ul>
                    <h5 style='color: #1976d2; margin-top: 15px; margin-bottom: 10px;'>Éléments mis en GRAS :</h5>
                    <ul style='color: #424242; margin-bottom: 0;'>
                        <li>✅ Nom complet de l'employé</li>
                        <li>✅ Date de naissance</li>
                        <li>✅ Lieu de naissance</li>
                        <li>✅ Poste de travail</li>
                        <li>✅ Date d'entrée</li>
                        <li>✅ Date de génération</li>
                        <li>✅ TOUTES les valeurs remplacées</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                
        else:
            st.markdown("""
            <div style='background: linear-gradient(135deg, #ff6b6b 0%, #ffa500 100%); 
                        padding: 20px; border-radius: 10px; margin: 20px 0;'>
                <h3 style='color: white; margin-bottom: 15px;'>⚠️ Template manquant</h3>
                <p style='color: #f8f9fa; margin: 0;'>
                    Le fichier template "<strong>attestation de travail.docx</strong>" n'a pas été trouvé.<br>
                    Veuillez vous assurer que le template est présent dans le dossier du projet.
                </p>
            </div>
            """, unsafe_allow_html=True)
    
    else:
        st.markdown("""
        <div style='background: linear-gradient(135deg, #ff6b6b 0%, #ffa500 100%); 
                    padding: 20px; border-radius: 10px; margin: 20px 0;'>
            <h3 style='color: white; margin-bottom: 15px;'>🔧 Installation requise</h3>
            <p style='color: #f8f9fa; margin-bottom: 15px;'>
                Pour utiliser le générateur d'attestations, vous devez installer la bibliothèque python-docx.
            </p>
            <p style='color: #f8f9fa; margin: 0;'>
                Exécutez dans votre terminal : <code style='background: rgba(0,0,0,0.3); padding: 3px 6px; border-radius: 3px;'>pip install python-docx</code>
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Section d'analytics organisationnelles avancées
    st.markdown("---")
    st.header("ANALYTICS ORGANISATIONNELLES")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Métriques Organisationnelles")
        
        # Création des métriques basées uniquement sur nos données
        our_metrics = {
            'Métrique': ['Âge Moyen', 'Taux de Cadres', 'Ratio Diversité', 'Ancienneté Moyenne', 'Employés Senior (55+)'],
            'Valeur': [
                f"{filtered_df['Age_calcule'].mean():.1f} ans" if 'Age_calcule' in filtered_df.columns else "N/A",
                f"{(filtered_df['CSP'].str.contains('Cadre|Manager|Directeur', case=False, na=False).sum() / len(filtered_df) * 100):.1f}%" if 'CSP' in filtered_df.columns else "N/A",
                f"{(min(filtered_df['Sexe'].value_counts()) / max(filtered_df['Sexe'].value_counts()) * 100):.1f}%" if 'Sexe' in filtered_df.columns else "N/A",
                f"{filtered_df['Anciennete_calculee'].mean():.1f} ans" if 'Anciennete_calculee' in filtered_df.columns else "N/A",
                f"{((filtered_df['Age_calcule'] >= 55).sum() / len(filtered_df) * 100):.1f}%" if 'Age_calcule' in filtered_df.columns else "N/A"
            ],
            'Observations': [
                'Basé sur dates de naissance',
                'Analyse des postes actuels', 
                'Équilibre hommes/femmes',
                'Depuis date d\'entrée',
                'Proche de la retraite'
            ]
        }
        
        metrics_df = pd.DataFrame(our_metrics)
        st.dataframe(metrics_df, use_container_width=True, hide_index=True)
    
    with col2:
        st.subheader("Indicateurs de Performance")
        
        # Affichage des métriques basées uniquement sur les données réelles
        if 'Age_calcule' in filtered_df.columns and 'Anciennete_calculee' in filtered_df.columns:
            metrics_data = {
                'Indicateur': ['Stabilité Équipe', 'Diversité Genre', 'Maturité Org.', 'Renouvellement'],
                'Score': [
                    f'{(filtered_df["Anciennete_calculee"] >= 5).sum() / len(filtered_df) * 100:.1f}%' if 'Anciennete_calculee' in filtered_df.columns else 'N/A',
                    f'{(min(filtered_df["Sexe"].value_counts()) / max(filtered_df["Sexe"].value_counts()) * 100):.1f}%' if 'Sexe' in filtered_df.columns else 'N/A',
                    f'{(filtered_df["CSP"].str.contains("Cadre|Manager", case=False, na=False).sum() / len(filtered_df) * 100):.1f}%' if 'CSP' in filtered_df.columns else 'N/A',
                    f'{(filtered_df["Age_calcule"] <= 35).sum() / len(filtered_df) * 100:.1f}%' if 'Age_calcule' in filtered_df.columns else 'N/A'
                ],
                'Description': [
                    'Employés avec 5+ ans d\'ancienneté',
                    'Équilibre hommes/femmes',
                    'Taux d\'encadrement',
                    'Employés de moins de 35 ans'
                ]
            }
            
            metrics_perf_df = pd.DataFrame(metrics_data)
            st.dataframe(metrics_perf_df, use_container_width=True, hide_index=True)
        
        # Graphique de répartition par direction (données réelles)
        if 'Direction' in filtered_df.columns:
            direction_counts = filtered_df['Direction'].value_counts().head(8)
            
            fig_directions = px.bar(
                x=direction_counts.values,
                y=direction_counts.index,
                orientation='h',
                title="Répartition par Direction",
                color=direction_counts.values,
                color_continuous_scale='viridis'
            )
            fig_directions.update_layout(
                xaxis_title="Nombre d'employés",
                yaxis_title="Direction",
                height=400
            )
            st.plotly_chart(fig_directions, use_container_width=True, key="directions_analytics_chart")
    
    st.markdown("---")
    st.header("DONNÉES DÉTAILLÉES & EXPORTS")
    
    # Onglets pour les données
    tab1, tab2, tab3 = st.tabs(["Liste des Employés", "Stats par Département", "Tableaux de Bord"])
    
    with tab1:
        st.subheader("Base de Données Employés Filtrée")
        
        # Options d'affichage
        col1, col2, col3 = st.columns(3)
        
        with col1:
            show_all = st.checkbox("Afficher toutes les colonnes")
        
        with col2:
            if st.button("Exporter en CSV"):
                csv = filtered_df.to_csv(index=False)
                st.download_button(
                    label="Télécharger CSV",
                    data=csv,
                    file_name=f"export_rh_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv"
                )
        
        with col3:
            st.metric("Nombre d'enregistrements", len(filtered_df))
        
        # Affichage du tableau
        if show_all:
            st.dataframe(filtered_df, use_container_width=True, height=400)
        else:
            columns_to_show = ['Matricule', 'Nom', 'Prenoms', 'Age_calcule', 'Sexe', 'Poste', 'Direction', 'Type de contrat', 'Anciennete_calculee']
            available_columns = [col for col in columns_to_show if col in filtered_df.columns]
            if available_columns:
                st.dataframe(filtered_df[available_columns], use_container_width=True, height=400)
    
    with tab2:
        st.subheader("Analyses par Département")
        
        if 'Déparetement' in filtered_df.columns:
            dept_stats = filtered_df.groupby('Déparetement').agg({
                'Matricule': 'count',
                'Age_calcule': ['mean', 'median'] if 'Age_calcule' in filtered_df.columns else 'count',
                'Anciennete_calculee': ['mean', 'median'] if 'Anciennete_calculee' in filtered_df.columns else 'count'
            }).round(1)
            
            # Aplatir les colonnes multi-niveaux
            if 'Age_calcule' in filtered_df.columns and 'Anciennete_calculee' in filtered_df.columns:
                dept_stats.columns = ['Effectif', 'Âge Moyen', 'Âge Médian', 'Ancienneté Moy.', 'Ancienneté Médiane']
            else:
                dept_stats.columns = ['Effectif']
            
            dept_stats = dept_stats.sort_values('Effectif', ascending=False)
            
            # Ajouter des calculs de ratios
            dept_stats['% du Total'] = (dept_stats['Effectif'] / len(filtered_df) * 100).round(1)
            
            st.dataframe(dept_stats, use_container_width=True)
            
            # Graphique des départements
            fig_dept = px.bar(
                x=dept_stats.index[:10],
                y=dept_stats['Effectif'][:10],
                title="Top 10 Départements par Effectif",
                color=dept_stats['Effectif'][:10],
                color_continuous_scale='viridis'
            )
            fig_dept.update_layout(xaxis_title="Département", yaxis_title="Effectif")
            st.plotly_chart(fig_dept, use_container_width=True, key="departements_chart")
        else:
            st.info("La colonne 'Département' n'est pas disponible dans les données.")
    
    with tab3:
        st.subheader("Tableaux de Bord Personnalisés")
        
        # Générateur de graphiques personnalisés
        col1, col2 = st.columns(2)
        
        with col1:
            x_axis = st.selectbox(
                "Axe X (Catégorie)",
                ['Direction', 'CSP', 'Sexe', 'Generation', 'Segment_Anciennete', 'Statut_Retraite']
            )
        
        with col2:
            chart_type = st.selectbox(
                "Type de graphique",
                ['Bar Chart', 'Pie Chart', 'Histogram']
            )
        
        if x_axis in filtered_df.columns:
            if chart_type == 'Bar Chart':
                counts = filtered_df[x_axis].value_counts()
                fig_custom = px.bar(x=counts.index, y=counts.values, 
                                  title=f"Répartition par {x_axis}")
                st.plotly_chart(fig_custom, use_container_width=True, key="custom_bar_chart")
            
            elif chart_type == 'Pie Chart':
                counts = filtered_df[x_axis].value_counts()
                fig_custom = px.pie(values=counts.values, names=counts.index,
                                  title=f"Répartition par {x_axis}")
                st.plotly_chart(fig_custom, use_container_width=True, key="custom_pie_chart")
    
    # Section Générateur d'Attestation de Travail
    st.markdown("---")
    st.header("GÉNÉRATEUR D'ATTESTATION DE TRAVAIL")
    
    # Vérifier si le template existe
    template_path = "attestation de travail.docx"
    
    if not DOCX_AVAILABLE:
        st.error("""
        ⚠️ **Fonctionnalité non disponible**
        
        Pour utiliser le générateur d'attestation, vous devez installer la bibliothèque python-docx :
        ```
        pip install python-docx
        ```
        """)
    elif not os.path.exists(template_path):
        st.error(f"""
        ❌ **Template introuvable**
        
        Le fichier '{template_path}' n'existe pas dans le dossier actuel.
        Veuillez vous assurer que le template PROMASIDOR soit présent dans le même dossier que ce script.
        """)
    else:
        st.success(f"✅ Template PROMASIDOR trouvé : {template_path}")
        
        st.markdown("""
        <div class="info-box">
            <h4>📋 Génération Automatique d'Attestations PROMASIDOR</h4>
            <p>Sélectionnez un employé pour générer automatiquement son attestation de travail personnalisée.</p>
            <p><em>Le document utilisera votre template officiel PROMASIDOR avec logo et formatage.</em></p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # Sélection de l'employé
            if len(df) > 0:
                # Créer une liste des employés avec nom complet
                if 'Nom' in df.columns and 'Prenoms' in df.columns:
                    df['Nom_Complet'] = df['Nom'].astype(str) + ' ' + df['Prenoms'].astype(str)
                    employee_options = df['Nom_Complet'].dropna().unique().tolist()
                    employee_options.sort()
                    
                    selected_employee = st.selectbox(
                        "🔍 Sélectionner un employé",
                        options=['-- Choisir un employé --'] + employee_options,
                        help="Tapez pour rechercher un employé"
                    )
                    
                    if selected_employee != '-- Choisir un employé --':
                        # Récupérer les données de l'employé sélectionné
                        employee_row = df[df['Nom_Complet'] == selected_employee].iloc[0]
                        
                        # Afficher les informations de l'employé
                        st.subheader(f"📊 Informations de {selected_employee}")
                        
                        col_info1, col_info2 = st.columns(2)
                        
                        with col_info1:
                            st.markdown(f"""
                            **Informations personnelles :**
                            - **Nom complet :** {employee_row.get('Nom', 'N/A')} {employee_row.get('Prenoms', 'N/A')}
                            - **Date de naissance :** {employee_row.get('Date de naissance', pd.NaT).strftime('%d/%m/%Y') if pd.notna(employee_row.get('Date de naissance')) else 'Non renseigné'}
                            - **Lieu de naissance :** {employee_row.get('Lieu de naissance', 'Non renseigné')}
                            - **Sexe :** {employee_row.get('Sexe', 'Non renseigné')}
                            """)
                        
                        with col_info2:
                            st.markdown(f"""
                            **Informations professionnelles :**
                            - **Poste :** {employee_row.get('Poste', 'Non renseigné')}
                            - **Direction :** {employee_row.get('Direction', 'Non renseigné')}
                            - **Date d'entrée :** {employee_row.get('DateEntree', pd.NaT).strftime('%d/%m/%Y') if pd.notna(employee_row.get('DateEntree')) else 'Non renseigné'}
                            - **Type de contrat :** {employee_row.get('Type de contrat', 'Non renseigné')}
                            - **Ancienneté :** {employee_row.get('Anciennete_calculee', 0):.1f} ans
                            """)
                        
                        # Bouton de génération
                        st.markdown("<br>", unsafe_allow_html=True)
                        
                        col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
                        with col_btn2:
                            if st.button("📄 Générer l'Attestation de Travail", type="primary", use_container_width=True):
                                with st.spinner("Génération de l'attestation en cours..."):
                                    # Générer l'attestation avec le template spécifique
                                    doc_buffer, error = generate_work_certificate(employee_row, template_path)
                                    
                                    if error:
                                        st.error(f"❌ {error}")
                                    else:
                                        st.success("✅ Attestation générée avec succès !")
                                        
                                        # Bouton de téléchargement
                                        filename = f"Attestation_{employee_row.get('Nom', '')}_{employee_row.get('Prenoms', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
                                        filename = filename.replace(' ', '_')
                                        
                                        st.download_button(
                                            label="💾 Télécharger l'Attestation",
                                            data=doc_buffer.getvalue(),
                                            file_name=filename,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True
                                        )
                                        
                                        st.markdown(f"""
                                        <div class="success-box">
                                            <h4>📋 Attestation Générée</h4>
                                            <p><strong>Employé :</strong> {selected_employee}</p>
                                            <p><strong>Date de génération :</strong> {datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
                                            <p><strong>Nom du fichier :</strong> {filename}</p>
                                        </div>
                                        """, unsafe_allow_html=True)
                else:
                    st.error("❌ Les colonnes 'Nom' et 'Prenoms' sont requises pour générer les attestations.")
            else:
                st.warning("⚠️ Aucune donnée d'employé disponible.")
        
        with col2:
            st.markdown("""
            <div class="highlight-box">
                <h4>ℹ️ Informations</h4>
                <p><strong>Template requis :</strong></p>
                <p>• "attestation de travail.docx"</p>
                <p><strong>Placeholders disponibles :</strong></p>
                <ul style="font-size: 0.8em;">
                    <li>[NOM_COMPLET]</li>
                    <li>[DATE_NAISSANCE]</li>
                    <li>[LIEU_NAISSANCE]</li>
                    <li>[POSTE]</li>
                    <li>[DATE_ENTREE]</li>
                    <li>[DATE_GENERATION]</li>
                </ul>
                <p style="font-size: 0.8em;"><em>Si le template n'existe pas, un modèle par défaut sera créé.</em></p>
            </div>
            """, unsafe_allow_html=True)
            
            # Statistiques rapides
            st.markdown("""
            <div class="info-box">
                <h4>📊 Statistiques</h4>
            </div>
            """, unsafe_allow_html=True)
            
            if len(df) > 0:
                total_employees = len(df)
                with_birthdate = df['Date de naissance'].notna().sum() if 'Date de naissance' in df.columns else 0
                with_entry_date = df['DateEntree'].notna().sum() if 'DateEntree' in df.columns else 0
                
                st.metric("Total employés", total_employees)
                st.metric("Avec date de naissance", f"{with_birthdate}/{total_employees}")
                st.metric("Avec date d'entrée", f"{with_entry_date}/{total_employees}")
    
    # Footer professionnel
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; padding: 20px; background-color: #f8f9fa; border-radius: 10px;'>
        <p style='color: #6c757d; margin: 0;'>
            <strong>Dashboard RH Executive</strong> | 
            <em>Powered by Advanced Analytics</em> | 
            <em>Dernière mise à jour: {}</em>
        </p>
        <p style='color: #6c757d; margin: 5px 0 0 0; font-size: 0.9em;'>
            <em>Tableau de bord généré automatiquement - Données actualisées en temps réel</em>
        </p>
    </div>
    """.format(datetime.now().strftime('%d/%m/%Y à %H:%M')), unsafe_allow_html=True)

if __name__ == "__main__":
    main()
